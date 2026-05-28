"""
Fleet Running Hours Monitor  v16
Architecture: Text-First Parser + Strict Native Type-Safe UI
"""
import streamlit as st
st.set_page_config(
    page_title="Fleet Running Hours Monitor",
    page_icon="⚓",
    layout="wide",
    initial_sidebar_state="collapsed",
)

import os
import re
import json
import shutil
import sqlite3
import hashlib
import tempfile
import subprocess
from pathlib import Path
from datetime import datetime
from typing import Any, Dict, List, Tuple

import pandas as pd

# ══════════════════════════════════════════════════════════════════
#  DESIGN
# ══════════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@400;500;600;700&family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');
:root{
  --bg:#071019; --bg2:#0c1623; --bg3:#111e30; --line:#1b2d44; --line2:#2a476b;
  --gold:#c99818; --gold2:#f0c45f; --red:#ef6666; --ora:#ffb04d; --grn:#69c07a; --blu:#6bb4ff;
  --t0:#ebf3ff; --t1:#a9bdd4; --t2:#71879f; --t3:#43576d; --r:14px;
  --ff:'Space Grotesk',sans-serif; --fi:'Inter',sans-serif; --fm:'JetBrains Mono',monospace;
}
html,body,[class*="css"]{background:var(--bg)!important;color:var(--t1)!important;font-family:var(--fi)!important}
.main,.main>div,.block-container{background:var(--bg)!important}
.block-container{max-width:100%!important;padding:1rem 1.5rem 3rem!important}
[data-testid="collapsedControl"],[data-testid="stSidebar"]{display:none!important}
.main::before{
 content:"";position:fixed;inset:0;pointer-events:none;z-index:0;
 background:
 radial-gradient(ellipse 70% 45% at 0% 0%, rgba(201,152,24,.08), transparent 60%),
 radial-gradient(ellipse 55% 35% at 100% 100%, rgba(107,180,255,.06), transparent 60%);
}
.block-container>*{position:relative;z-index:1}
.hero-k{font-size:.66rem;letter-spacing:.24em;text-transform:uppercase;color:var(--gold2);font-weight:700}
.hero-h{font-family:var(--ff);font-size:1.9rem;font-weight:700;color:var(--t0);letter-spacing:-.04em;line-height:1.06;margin-top:.2rem}
.hero-s{color:var(--t1);font-size:.92rem;line-height:1.65;margin-top:.45rem;max-width:1100px}
.hero-rule{height:1px;margin:.8rem 0 1rem;background:linear-gradient(90deg,var(--gold2),var(--line),transparent)}
.metric-grid{display:grid;grid-template-columns:repeat(8,1fr);gap:.75rem;margin:1rem 0}
.metric{
 background:linear-gradient(180deg,var(--bg3),var(--bg2));
 border:1px solid var(--line);border-radius:var(--r);padding:.85rem .95rem;position:relative
}
.metric::before{content:"";position:absolute;left:0;right:0;top:0;height:2px;background:linear-gradient(90deg,var(--gold),transparent 75%)}
.metric.r::before{background:linear-gradient(90deg,var(--red),transparent 75%)}
.metric.o::before{background:linear-gradient(90deg,var(--ora),transparent 75%)}
.metric.g::before{background:linear-gradient(90deg,var(--grn),transparent 75%)}
.metric.b::before{background:linear-gradient(90deg,var(--blu),transparent 75%)}
.metric-v{font-family:var(--ff);font-size:1.45rem;font-weight:700;color:var(--t0);line-height:1.05;letter-spacing:-.04em}
.metric-l{color:var(--t3);font-size:.58rem;text-transform:uppercase;letter-spacing:.16em;margin-top:.3rem}
.bar{display:flex;align-items:center;gap:.75rem;margin:1.35rem 0 .8rem}
.bar-i{width:34px;height:34px;border-radius:10px;display:flex;align-items:center;justify-content:center;background:rgba(201,152,24,.1);border:1px solid rgba(201,152,24,.2)}
.bar-t{font-family:var(--ff);font-size:1rem;font-weight:700;color:var(--t0)}
.bar-s{font-family:var(--fm);font-size:.6rem;color:var(--t3);margin-top:2px}
.bar-r{flex:1;height:1px;background:linear-gradient(90deg,var(--line),transparent)}
.bar-b{font-family:var(--fm);font-size:.6rem;color:var(--t2);border:1px solid var(--line2);background:var(--bg3);border-radius:999px;padding:4px 9px}
.banner{border-radius:var(--r);padding:.9rem 1rem;border:1px solid var(--line);margin:.85rem 0;line-height:1.55}
.banner.ok{background:rgba(105,192,122,.1);border-color:rgba(105,192,122,.28);color:#e6faea}
.banner.warn{background:rgba(255,176,77,.1);border-color:rgba(255,176,77,.28);color:#ffe9cd}
.banner.err{background:rgba(239,102,102,.1);border-color:rgba(239,102,102,.28);color:#ffdede}
[data-testid="stFileUploadDropzone"]{
 background:rgba(201,152,24,.04)!important;border:1.5px dashed rgba(201,152,24,.5)!important;
 border-radius:16px!important;padding:2rem 1.25rem!important
}
[data-testid="stFileUploadDropzone"]:hover{border-color:var(--gold2)!important;background:rgba(201,152,24,.07)!important}
[data-testid="stDataFrame"]{border:1px solid var(--line)!important;border-radius:var(--r)!important;overflow:auto!important}
.dvn-scroller{background:var(--bg2)!important}
.streamlit-expanderHeader{background:var(--bg3)!important;border:1px solid var(--line)!important;border-radius:12px!important;color:var(--t0)!important}
.streamlit-expanderContent{background:var(--bg2)!important;border:1px solid var(--line)!important;border-top:none!important;border-radius:0 0 12px 12px!important}
.stButton>button{
 background:linear-gradient(135deg,var(--gold2),var(--gold))!important;color:#08111b!important;border:none!important;
 border-radius:11px!important;padding:.72rem 1rem!important;font-weight:800!important;text-transform:uppercase!important
}
.stDownloadButton>button{
 background:var(--bg3)!important;color:var(--t0)!important;border:1px solid var(--line2)!important;border-radius:11px!important
}
div[data-baseweb="select"] > div,.stTextInput>div>div>input{
 background:var(--bg3)!important;color:var(--t0)!important;border:1px solid var(--line2)!important;border-radius:10px!important
}
.stSelectbox label,.stMultiSelect label,.stCheckbox label,.stRadio label{
 color:var(--t3)!important;font-size:.66rem!important;text-transform:uppercase!important;letter-spacing:.12em!important
}
@media(max-width:1280px){.metric-grid{grid-template-columns:repeat(4,1fr)}}
@media(max-width:860px){.metric-grid{grid-template-columns:repeat(2,1fr)}}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════
#  CONSTANTS
# ══════════════════════════════════════════════════════════════════
DB_PATH = "tec004_reports.db"
PARSER_VERSION = "16.0_safe_render"

STATUS_ORDER = {'OVERDUE': 0, 'HIGH PRIORITY': 1, 'OK': 2, 'NO DATA': 3}
STATUS_MAP = {
    'OVERDUE': '🔴 OVERDUE',
    'HIGH PRIORITY': '🟠 HIGH PRIORITY',
    'OK': '🟢 OK',
    'NO DATA': '🔵 NO DATA'
}

# ══════════════════════════════════════════════════════════════════
#  DB
# ══════════════════════════════════════════════════════════════════
def get_conn():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("""
    CREATE TABLE IF NOT EXISTS reports(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        vessel_name TEXT NOT NULL,
        report_date TEXT,
        filename TEXT NOT NULL,
        file_hash TEXT NOT NULL,
        parser_version TEXT NOT NULL,
        me_total_hrs REAL,
        me_this_month REAL,
        raw_json TEXT,
        created_at TEXT NOT NULL,
        UNIQUE(vessel_name, report_date, file_hash)
    )
    """)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS parsed_rows(
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        report_id INTEGER NOT NULL,
        category TEXT,
        engine_label TEXT,
        unit TEXT,
        description TEXT,
        periodicity REAL,
        last_oh_date TEXT,
        hrs_since REAL,
        pct_used REAL,
        status TEXT,
        parser_source TEXT,
        source_ref TEXT,
        created_at TEXT NOT NULL
    )
    """)
    conn.commit()
    conn.close()

init_db()

# ══════════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════════
def now_iso():
    return datetime.utcnow().isoformat()

def md5_bytes(b: bytes) -> str:
    return hashlib.md5(b).hexdigest()

def fl(txt: Any) -> str:
    raw = str(txt or "").replace('\x07', '').replace('\xa0', ' ').replace('\t', ' ')
    raw = raw.replace('\r', '\n').replace('\x0b', '\n')
    parts = [re.sub(r'\s+', ' ', p).strip() for p in raw.split('\n')]
    parts = [p for p in parts if p]
    return parts[0] if parts else ''

def clean_name(txt: Any) -> str:
    t = fl(txt)
    t = re.sub(r'(?i)^MV\s+', '', t)
    t = re.sub(r'(?i)Page\s*\d+\s*of\s*\d+', '', t)
    t = re.sub(r'\s+', ' ', t)
    return t.strip(" :-")

def parse_number(txt: Any) -> float:
    s = fl(txt).upper().replace('[', '').replace(']', '')
    if not s or s in ('-', 'N/A', 'NA', 'CENTRAL', 'COOLER'):
        return 0.0
    if any(w in s for w in ('MONTH', 'YEAR', 'WEEK', 'DAY', 'OBSERVATION')):
        return 0.0
    m = re.search(r'\d[\d,\.]*', s)
    if not m:
        return 0.0
    block = m.group()
    sep = max(block.rfind('.'), block.rfind(','))
    if sep > 0 and len(block) - sep == 4:
        block = re.sub(r'[,\.]', '', block)
    elif sep > 0:
        block = re.sub(r'[,\.]', '', block[:sep])
    else:
        block = re.sub(r'[,\.]', '', block)
    try:
        return float(block)
    except Exception:
        return 0.0

def parse_date(txt: Any) -> str:
    s = fl(txt).replace('[', '').replace(']', '').strip()
    if not s or s in ('-', '1', '2', 'N/A', 'NA', 'Central', 'cooler', 'CENTRAL', 'COOLER'):
        return ''
    if re.fullmatch(r'\d+', s):
        return ''
    if re.match(r'^\d{1,2}[/-]\d{1,2}[/-]\d{2,4}$', s):
        dd, mm, yy = re.split(r'[/-]', s)
        try:
            if not (1 <= int(dd) <= 31 and 1 <= int(mm) <= 12):
                return ''
        except Exception:
            return ''
        return f"{int(dd):02d}/{int(mm):02d}/{str(yy)[-2:]}"
    if re.match(r'^\d{1,2}\s+[A-Za-z\.]+\s+\d{2,4}$', s):
        return s
    return s if re.search(r'[A-Za-z/]', s) else ''

def is_component_name(name: str) -> bool:
    u = fl(name).upper()
    if not u or len(u) < 2:
        return False
    bad = (
        'DESCRIPTION','PERIODICTLY','PERIODICITY','MAIN ENGINE','AUX. ENGINE','AUX. ENGINE MAKER / TYPE',
        'TOTAL HOURS','HOURS THIS MONTH','SERIAL NR','RUNNING HOURS MONTHLY REPORT',
        'DATE OF LAST O/H','RUNNING HOURS SINCE LAST O/H','REMARKS',
        'TURBOCHARGER','A/C & REFR. COMPRESSORS','MAIN AIR COMPRESSORS','COOLERS'
    )
    if any(b in u for b in bad):
        return False
    if re.fullmatch(r'[\d./ ,:\-\(\)\[\]]+', u):
        return False
    return bool(re.search(r'[A-Z]', u))

def status_of(hrs: float, period: float) -> str:
    if hrs <= 0 or period <= 0:
        return 'NO DATA'
    r = hrs / period
    if r >= 1.0:
        return 'OVERDUE'
    if r >= 0.8:
        return 'HIGH PRIORITY'
    return 'OK'

def pct_used(hrs: float, period: float) -> float:
    return round(hrs / period, 4) if hrs and period else 0.0

def make_row(cat, eng, unit, name, period, date, hrs, parser_source='text', source_ref=''):
    return {
        'category': cat,
        'engine_label': eng,
        'unit': unit,
        'description': name,
        'periodicity': period,
        'last_oh_date': date,
        'hrs_since': hrs,
        'pct_used': pct_used(hrs, period),
        'status': status_of(hrs, period),
        'parser_source': parser_source,
        'source_ref': source_ref,
    }

# ══════════════════════════════════════════════════════════════════
#  DOC CONVERSION
# ══════════════════════════════════════════════════════════════════
def convert_doc_to_docx(raw: bytes) -> bytes:
    soffice = shutil.which("soffice") or "/usr/bin/soffice"
    if not os.path.isfile(soffice):
        raise RuntimeError("LibreOffice not found. Install libreoffice in the runtime.")
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as t:
        t.write(raw)
        src = t.name
    outdir = tempfile.mkdtemp(prefix="lo_")
    out = os.path.join(outdir, Path(src).stem + ".docx")
    profile = f"file:///tmp/lo_{os.getpid()}_{os.urandom(4).hex()}"
    try:
        r = subprocess.run(
            [soffice, "--headless", "--norestore", "--nofirststartwizard",
             f"-env:UserInstallation={profile}", "--convert-to", "docx", src, "--outdir", outdir],
            capture_output=True, timeout=120
        )
        if not os.path.exists(out):
            raise RuntimeError(r.stderr.decode("utf-8", "ignore")[:500] or r.stdout.decode("utf-8", "ignore")[:500])
        with open(out, "rb") as f:
            return f.read()
    finally:
        for p in [src, out]:
            try:
                if os.path.exists(p):
                    os.unlink(p)
            except Exception:
                pass
        shutil.rmtree(outdir, ignore_errors=True)

# ══════════════════════════════════════════════════════════════════
#  EXTRACTION
# ══════════════════════════════════════════════════════════════════
def table_to_grid(table) -> List[List[str]]:
    grid = []
    max_cols = 0
    for row in table.rows:
        vals = []
        for cell in row.cells:
            vals.append(fl(cell.text))
        max_cols = max(max_cols, len(vals))
        grid.append(vals)
    for row in grid:
        while len(row) < max_cols:
            row.append('')
    return grid

def extract_doc_lines(doc) -> List[str]:
    lines = []
    for p in doc.paragraphs:
        t = fl(p.text)
        if t:
            lines.append(t)
    for table in doc.tables:
        grid = table_to_grid(table)
        for row in grid:
            for cell in row:
                t = fl(cell)
                if t:
                    lines.append(t)
    return lines

def extract_vessel_and_date(lines: List[str]) -> Tuple[str, str]:
    vessel = 'UNKNOWN'
    report_date = ''
    for line in lines[:60]:
        if 'VESSEL' in line.upper() and 'DATE' in line.upper():
            if m := re.search(r"Vessel[\u2019\u2018']?s?\s+Name\s*:\s*(?:MV\s+)?(.+?)\s+Date\s*:\s*(.+)$", line, re.I):
                vessel = clean_name(m.group(1))
                report_date = parse_date(m.group(2))
                break
    return vessel, report_date

def extract_me_totals(lines: List[str]) -> Tuple[float, float]:
    me_total = 0.0
    me_month = 0.0
    joined = "\n".join(lines[:120])
    if m := re.search(r'Total Running Hours\s*[:ǀ|]?\s*([\d,]+)', joined, re.I):
        me_total = parse_number(m.group(1))
    if m := re.search(r'This Month\s*[:]?\s*([\d,]+)', joined, re.I):
        me_month = parse_number(m.group(1))
    return me_total, me_month

def section_between(lines: List[str], start_markers: List[str], end_markers: List[str]) -> List[str]:
    start = None
    end = len(lines)
    for i, line in enumerate(lines):
        up = line.upper()
        if start is None and any(m in up for m in start_markers):
            start = i
            continue
        if start is not None and any(m in up for m in end_markers):
            end = i
            break
    return lines[start:end] if start is not None else []

def normalize_component_token(line: str) -> str:
    return clean_name(line).replace('## ', '').strip()

def parse_me_text(lines: List[str]) -> Tuple[List[Dict[str, Any]], List[str]]:
    issues = []
    out = []
    sec = section_between(
        lines,
        ['MAIN ENGINE', 'CYL. NO.1'],
        ['NOTE 1', 'TURBOCHARGER', 'AUX. ENGINE MAKER / TYPE']
    )
    if not sec:
        issues.append("Main Engine section not found.")
        return out, issues

    stream = [fl(x) for x in sec if fl(x)]
    stream = [x for x in stream if x.upper() not in ('CYL. NO.1', 'CYL. NO.2', 'CYL. NO.3', 'CYL. NO.4', 'CYL. NO.5', 'CYL. NO.6', 'CYL. NO.7')]

    i = 0
    while i < len(stream):
        name = normalize_component_token(stream[i])
        if is_component_name(name):
            period = parse_number(stream[i + 1]) if i + 1 < len(stream) else 0.0
            marker1 = fl(stream[i + 2]) if i + 2 < len(stream) else ''
            if marker1 == '1':
                dates = []
                j = i + 3
                while j < len(stream) and fl(stream[j]) != '2' and len(dates) < 7:
                    token = fl(stream[j])
                    if token and token.upper() != 'REMARKS':
                        dates.append(token)
                    j += 1

                if j < len(stream) and fl(stream[j]) == '2':
                    j += 1
                    hrs = []
                    while j < len(stream) and len(hrs) < max(1, len(dates)):
                        token = fl(stream[j])
                        if is_component_name(token):
                            break
                        hrs.append(token)
                        j += 1

                    for idx in range(max(len(dates), len(hrs))):
                        d = parse_date(dates[idx]) if idx < len(dates) else ''
                        h = parse_number(hrs[idx]) if idx < len(hrs) else 0.0
                        if d or h > 0:
                            out.append(make_row(
                                'MAIN_ENGINE', 'ME', f'Cyl {idx+1}', name, period, d, h,
                                parser_source='text', source_ref=f"ME:{name}"
                            ))
                    i = j
                    continue
        i += 1

    if not out:
        issues.append("Main Engine parse returned zero rows.")
    return out, issues

def parse_aux_text(lines: List[str]) -> Tuple[List[Dict[str, Any]], List[str]]:
    issues = []
    out = []
    sec = section_between(
        lines,
        ['AUX. ENGINE MAKER / TYPE', 'AUX. ENGINE NO.1'],
        ['DESCRIPTION', 'D/G NO1', 'D/G NO.1']
    )

    if not sec:
        sec = section_between(
            lines,
            ['AUX. ENGINE MAKER / TYPE', 'AUX. ENGINE NO.1'],
            ['1ST COPY TO BE RETAINED', 'REMARKS']
        )

    if not sec:
        issues.append("Auxiliary Engine section not found.")
        return out, issues

    stream = [fl(x) for x in sec if fl(x)]

    start_idx = -1
    for i in range(len(stream) - 3):
        if stream[i].upper() == 'DESCRIPTION' and 'PERIODICTLY' in stream[i + 1].upper() and stream[i + 2] == '1' and stream[i + 3] == '2':
            start_idx = i + 4
            break

    if start_idx == -1:
        for i in range(len(stream)):
            if stream[i].upper() == 'DESCRIPTION':
                start_idx = i + 1
                break

    if start_idx == -1:
        issues.append("Auxiliary Engine data header not found.")
        return out, issues

    data = stream[start_idx:]

    i = 0
    while i < len(data):
        name = normalize_component_token(data[i])
        if is_component_name(name):
            period = parse_number(data[i + 1]) if i + 1 < len(data) else 0.0
            marker1 = fl(data[i + 2]) if i + 2 < len(data) else ''
            if marker1 == '1':
                date_val = parse_date(data[i + 3]) if i + 3 < len(data) else ''
                marker2 = fl(data[i + 4]) if i + 4 < len(data) else ''
                hrs_val = parse_number(data[i + 5]) if i + 5 < len(data) else 0.0

                if marker2 == '2' and (date_val or hrs_val > 0):
                    out.append(make_row(
                        'AUX_ENGINE', 'AUX-1', 'General', name, period, date_val, hrs_val,
                        parser_source='text', source_ref=f"AUX:{name}"
                    ))
                    i += 6
                    continue
        i += 1

    if not out:
        issues.append("Auxiliary Engine parse returned zero rows.")
    return out, issues

def parse_other_equipment_text(lines: List[str]) -> List[Dict[str, Any]]:
    out = []
    sec = section_between(
        lines,
        ['TURBOCHARGER', 'A/C & REFR. COMPRESSORS'],
        ['AUX. ENGINE MAKER / TYPE']
    )
    stream = [fl(x) for x in sec if fl(x)]
    
    i = 0
    while i < len(stream) - 1:
        name = normalize_component_token(stream[i])
        if is_component_name(name):
            dt = ''
            hrs = 0.0
            offset = 1
            # Dynamic lookahead for robust OE extraction
            for offset in range(1, min(5, len(stream) - i)):
                tok = stream[i + offset]
                if not dt and ('/' in tok or '-' in tok or re.search(r'[A-Za-z]{3}\s+\d{2,4}', tok)):
                    dt = parse_date(tok)
                if hrs == 0.0 and re.match(r'^\d[\d,\.]*$', tok) and tok not in ('1', '2'):
                    hrs = parse_number(tok)
            
            if dt or hrs > 0:
                out.append({
                    'section': 'Other Equipment',
                    'description': name,
                    'last_date': dt if dt else '—',
                    'run_hrs': int(hrs) if hrs > 0 else '',
                })
                i += offset
                continue
        i += 1
        
    seen = set()
    unique_out = []
    for item in out:
        key = (item['section'], item['description'])
        if key not in seen:
            seen.add(key)
            unique_out.append(item)
            
    return unique_out

def parse_doc_bytes(docx_bytes: bytes, filename: str = '') -> Dict[str, Any]:
    from docx import Document

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as t:
        t.write(docx_bytes)
        tp = t.name
    try:
        doc = Document(tp)
    except Exception as e:
        raise ValueError(f"Cannot open converted DOCX: {e}")
    finally:
        try:
            os.unlink(tp)
        except Exception:
            pass

    if not doc.tables:
        raise ValueError("No tables found — verify this is a TEC-004 report.")

    raw_lines = extract_doc_lines(doc)
    vessel_name, report_date = extract_vessel_and_date(raw_lines)
    me_total, me_month = extract_me_totals(raw_lines)

    me_rows, me_issues = parse_me_text(raw_lines)
    aux_rows, aux_issues = parse_aux_text(raw_lines)
    oe_rows = parse_other_equipment_text(raw_lines)

    issues = []
    if vessel_name == 'UNKNOWN':
        issues.append({'severity': 'warning', 'message': 'Could not extract vessel name.'})
    if not report_date:
        issues.append({'severity': 'warning', 'message': 'Could not extract report date.'})
    issues += [{'severity': 'warning', 'message': x} for x in me_issues]
    issues += [{'severity': 'warning', 'message': x} for x in aux_issues]

    components = me_rows + aux_rows
    if not components:
        issues.append({'severity': 'error', 'message': 'No ME/AUX rows extracted from text stream.'})

    debug_tables = []
    for idx, table in enumerate(doc.tables):
        grid = table_to_grid(table)
        debug_tables.append({
            'table_index': idx,
            'rows': len(grid),
            'cols': max(len(r) for r in grid) if grid else 0,
            'sample': grid[:12]
        })

    return {
        'vessel_name': vessel_name,
        'report_date': report_date,
        'me_total_hrs': int(me_total) if me_total else 0,
        'me_this_month': int(me_month) if me_month else 0,
        'components': components,
        'me_comps': me_rows,
        'aux_comps': aux_rows,
        'other_equipment': oe_rows,
        'warnings': [x['message'] for x in issues if x['severity'] == 'warning'],
        'issues': issues,
        'debug': {
            'raw_lines_preview': raw_lines[:250],
            'table_debug': debug_tables,
            'me_rows_total': len(me_rows),
            'aux_rows_total': len(aux_rows),
        },
        'uploaded_at': now_iso(),
        'filename': filename,
    }

# ══════════════════════════════════════════════════════════════════
#  SAVE / LOAD
# ══════════════════════════════════════════════════════════════════
def report_exists(vessel_name: str, report_date: str, file_hash: str) -> bool:
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("SELECT id FROM reports WHERE vessel_name=? AND report_date=? AND file_hash=?", (vessel_name, report_date, file_hash))
    exists = cur.fetchone() is not None
    conn.close()
    return exists

def save_report(parsed: Dict[str, Any]) -> Tuple[bool, str]:
    conn = get_conn()
    try:
        if report_exists(parsed['vessel_name'], parsed['report_date'], parsed['file_hash']):
            return False, "This report is already saved."
        cur = conn.cursor()
        cur.execute("""
        INSERT INTO reports(
            vessel_name, report_date, filename, file_hash, parser_version,
            me_total_hrs, me_this_month, raw_json, created_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            parsed['vessel_name'], parsed['report_date'], parsed['filename'], parsed['file_hash'],
            PARSER_VERSION, parsed['me_total_hrs'], parsed['me_this_month'],
            json.dumps(parsed, ensure_ascii=False, default=str), now_iso()
        ))
        report_id = cur.lastrowid

        for r in parsed['components']:
            cur.execute("""
            INSERT INTO parsed_rows(
                report_id, category, engine_label, unit, description,
                periodicity, last_oh_date, hrs_since, pct_used, status,
                parser_source, source_ref, created_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                report_id, r['category'], r['engine_label'], r['unit'], r['description'],
                r['periodicity'], r['last_oh_date'], r['hrs_since'], r['pct_used'], r['status'],
                r['parser_source'], r['source_ref'], now_iso()
            ))
        conn.commit()
        return True, f"Saved successfully. Report ID {report_id}."
    except Exception as e:
        conn.rollback()
        return False, f"Save failed: {e}"
    finally:
        conn.close()

def load_recent_reports(limit=30) -> pd.DataFrame:
    conn = get_conn()
    try:
        df = pd.read_sql_query(f"""
        SELECT r.id, r.vessel_name, r.report_date, r.filename, r.parser_version,
               r.me_total_hrs, r.me_this_month, r.created_at,
               COUNT(pr.id) AS parsed_rows
        FROM reports r
        LEFT JOIN parsed_rows pr ON pr.report_id = r.id
        GROUP BY r.id
        ORDER BY r.id DESC
        LIMIT {int(limit)}
        """, conn)
        return df
    finally:
        conn.close()

# ══════════════════════════════════════════════════════════════════
#  DATAFRAME BUILDERS (CRASH-PROOF TYPE SAFETY)
# ══════════════════════════════════════════════════════════════════
def build_component_df(records: List[Dict[str, Any]], mode='matrix') -> pd.DataFrame:
    if not records:
        return pd.DataFrame()
    
    df = pd.DataFrame(records)
    df['_s'] = df['status'].map(lambda x: STATUS_ORDER.get(str(x), 9))
    df['_p'] = pd.to_numeric(df['pct_used'], errors='coerce').fillna(0.0)
    
    if mode == 'matrix':
        df['_k1'] = df['description'].astype(str).str.upper()
        df['_k2'] = df['engine_label'].astype(str)
        df = df.sort_values(['_k1','_k2','unit']).reset_index(drop=True)
    else:
        df = df.sort_values(['_s','_p'], ascending=[True, False]).reset_index(drop=True)

    out = pd.DataFrame()
    out['Status'] = df['status'].map(lambda x: STATUS_MAP.get(str(x), '🔵 NO DATA'))
    out['Component'] = df['description'].astype(str)
    out['Engine'] = df['engine_label'].astype(str)
    out['Unit'] = df['unit'].astype(str)
    
    # Cast to strict string to completely prevent PyArrow NumberColumn exceptions
    out['Periodicity'] = [f"{int(float(x))}" if pd.notna(x) and float(x) > 0 else "—" for x in df['periodicity']]
    out['Last O/H'] = [str(x) if x else '—' for x in df['last_oh_date']]
    out['Hrs Since'] = [f"{int(float(x))}" if pd.notna(x) and float(x) > 0 else "—" for x in df['hrs_since']]
    
    # Strict float ensuring zero NaN for the ProgressColumn
    out['Used %'] = [round(float(x)*100, 1) if pd.notna(x) and float(x) > 0 else 0.0 for x in df['pct_used']]
    
    return out

def build_issue_df(issues: List[Dict[str, Any]]) -> pd.DataFrame:
    if not issues:
        return pd.DataFrame(columns=['Severity','Message'])
    return pd.DataFrame(issues)[['severity','message']].rename(columns={'severity':'Severity','message':'Message'})

# ══════════════════════════════════════════════════════════════════
#  SESSION
# ══════════════════════════════════════════════════════════════════
if 'parsed_reports' not in st.session_state:
    st.session_state.parsed_reports = []
if 'active_report_hash' not in st.session_state:
    st.session_state.active_report_hash = None
if 'save_feedback' not in st.session_state:
    st.session_state.save_feedback = {}

# ══════════════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════════════
st.markdown("""
<div class="hero-k">Running Hours Management System</div>
<div class="hero-h">TEC‑004 Master Engine Build</div>
<div class="hero-s">
This framework enforces strict type-safe rendering across all metrics. The UI has been fully decoupled from Pandas styling loops, ensuring Streamlit matrices populate flawlessly without triggering internal PyArrow serialization crashes.
</div>
<div class="hero-rule"></div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════
#  UPLOAD
# ══════════════════════════════════════════════════════════════════
with st.expander("Upload TEC-004 Report", expanded=True):
    uploaded = st.file_uploader("Drop .doc file here", type=['doc'], key='tec004_file')
    if uploaded is not None:
        raw = uploaded.getvalue()
        fh = md5_bytes(raw)
        current = next((r for r in st.session_state.parsed_reports if r.get('file_hash') == fh), None)
        if current is None:
            with st.spinner("Converting .doc → .docx via LibreOffice…"):
                docx = convert_doc_to_docx(raw)
            with st.spinner("Parsing TEC‑004 from text stream…"):
                parsed = parse_doc_bytes(docx, filename=uploaded.name)
                parsed['file_hash'] = fh
                st.session_state.parsed_reports = [parsed]
                st.session_state.active_report_hash = fh

reports = st.session_state.parsed_reports
if not reports:
    st.info("Upload a TEC‑004 .doc report to parse it.")
    st.stop()

active = reports[0]
me = active['me_comps']
aux = active['aux_comps']
oe = active['other_equipment']
all_ = active['components']
issues = active['issues']
debug = active['debug']

n_od = sum(1 for c in all_ if c['status'] == 'OVERDUE')
n_hp = sum(1 for c in all_ if c['status'] == 'HIGH PRIORITY')
n_ok = sum(1 for c in all_ if c['status'] == 'OK')
n_err = sum(1 for x in issues if x['severity'] == 'error')
n_warn = sum(1 for x in issues if x['severity'] == 'warning')

st.markdown(f"""
<div class="metric-grid">
  <div class="metric b"><div class="metric-v">{active['vessel_name']}</div><div class="metric-l">Vessel</div></div>
  <div class="metric"><div class="metric-v">{active['report_date'] or '—'}</div><div class="metric-l">Report Date</div></div>
  <div class="metric b"><div class="metric-v">{active['me_total_hrs']:,}</div><div class="metric-l">ME Total Hrs</div></div>
  <div class="metric b"><div class="metric-v">{active['me_this_month']:,}</div><div class="metric-l">ME This Month</div></div>
  <div class="metric"><div class="metric-v">{len(me)}</div><div class="metric-l">ME Rows</div></div>
  <div class="metric"><div class="metric-v">{len(aux)}</div><div class="metric-l">AUX Rows</div></div>
  <div class="metric o"><div class="metric-v">{n_hp}</div><div class="metric-l">High Priority</div></div>
  <div class="metric r"><div class="metric-v">{n_od}</div><div class="metric-l">Overdue</div></div>
</div>
""", unsafe_allow_html=True)

if n_err:
    st.markdown(f'<div class="banner err">{n_err} parser errors detected.</div>', unsafe_allow_html=True)
elif n_warn:
    st.markdown(f'<div class="banner warn">{n_warn} parser warnings detected.</div>', unsafe_allow_html=True)
else:
    st.markdown('<div class="banner ok">Report parsed successfully with no warnings.</div>', unsafe_allow_html=True)

a1, a2, a3 = st.columns([1.3, 1.3, 3.4])
with a1:
    review_complete = st.checkbox("Review complete", value=False)
with a2:
    if st.button("Save active report", disabled=not review_complete):
        ok, msg = save_report(active)
        st.session_state.save_feedback[active['file_hash']] = (ok, msg)
with a3:
    export_df = build_component_df(all_, mode='matrix')
    st.download_button(
        "Download parsed CSV",
        export_df.to_csv(index=False).encode('utf-8'),
        file_name=f"{Path(active['filename']).stem}_parsed.csv",
        mime='text/csv'
    )

fb = st.session_state.save_feedback.get(active['file_hash'])
if fb:
    if fb[0]:
        st.success(fb[1])
    else:
        st.warning(fb[1])

tab1, tab2, tab3, tab4, tab5 = st.tabs(["Main Engine", "Auxiliary Engine", "Other Equipment", "Parse Issues", "Parser Debug"])

with tab1:
    st.markdown(f'<div class="bar"><div class="bar-i">⚙</div><div><div class="bar-t">Main Engine</div><div class="bar-s">paired row stream under cylinder sequence</div></div><div class="bar-r"></div><div class="bar-b">{len(me)} rows</div></div>', unsafe_allow_html=True)
    mf1, mf2, mf3 = st.columns([2, 2, 3])
    with mf1:
        me_comp = st.selectbox('Component', ['All'] + sorted({c['description'] for c in me}), key='me_comp')
    with mf2:
        me_status = st.selectbox('Status', ['All','Overdue only','High Priority +','OK only'], key='me_status')
    with mf3:
        me_sort = st.radio('Sort', ['Component → Cylinder', 'Priority → % Used'], horizontal=True, key='me_sort')

    me_view = me[:]
    if me_comp != 'All':
        me_view = [c for c in me_view if c['description'] == me_comp]
    if me_status == 'Overdue only':
        me_view = [c for c in me_view if c['status'] == 'OVERDUE']
    elif me_status == 'High Priority +':
        me_view = [c for c in me_view if c['status'] in ('OVERDUE', 'HIGH PRIORITY')]
    elif me_status == 'OK only':
        me_view = [c for c in me_view if c['status'] == 'OK']

    me_df = build_component_df(me_view, mode='matrix' if 'Component' in me_sort else 'priority')
    if me_df.empty:
        st.info("No data matches the selected filter.")
    else:
        # STRICT CONFIG: Numbers explicitly passed as TextColumn to bypass internal PyArrow crashes
        st.dataframe(
            me_df, use_container_width=True, hide_index=True,
            height=min(820, 38 * len(me_df) + 44),
            column_config={
                'Status': st.column_config.TextColumn("Status", width="small"),
                'Periodicity': st.column_config.TextColumn("Periodicity", width="small"),
                'Hrs Since': st.column_config.TextColumn("Hrs Since", width="small"),
                'Used %': st.column_config.ProgressColumn('Used %', min_value=0.0, max_value=150.0, format='%.1f%%', width="small")
            }
        )

with tab2:
    st.markdown(f'<div class="bar"><div class="bar-i">🔩</div><div><div class="bar-t">Auxiliary Engine</div><div class="bar-s">single component-pair stream from sample file</div></div><div class="bar-r"></div><div class="bar-b">{len(aux)} rows</div></div>', unsafe_allow_html=True)
    af1, af2 = st.columns([2, 2])
    with af1:
        ax_comp = st.selectbox('Component', ['All'] + sorted({c['description'] for c in aux}), key='ax_comp')
    with af2:
        ax_status = st.selectbox('Status', ['All','Overdue only','High Priority +','OK only'], key='ax_status')

    ax_view = aux[:]
    if ax_comp != 'All':
        ax_view = [c for c in ax_view if c['description'] == ax_comp]
    if ax_status == 'Overdue only':
        ax_view = [c for c in ax_view if c['status'] == 'OVERDUE']
    elif ax_status == 'High Priority +':
        ax_view = [c for c in ax_view if c['status'] in ('OVERDUE', 'HIGH PRIORITY')]
    elif ax_status == 'OK only':
        ax_view = [c for c in ax_view if c['status'] == 'OK']

    ax_df = build_component_df(ax_view)
    if ax_df.empty:
        st.info("No data matches the selected filter.")
    else:
        st.dataframe(
            ax_df, use_container_width=True, hide_index=True,
            height=min(820, 38 * len(ax_df) + 44),
            column_config={
                'Status': st.column_config.TextColumn("Status", width="small"),
                'Periodicity': st.column_config.TextColumn("Periodicity", width="small"),
                'Hrs Since': st.column_config.TextColumn("Hrs Since", width="small"),
                'Used %': st.column_config.ProgressColumn('Used %', min_value=0.0, max_value=150.0, format='%.1f%%', width="small")
            }
        )

with tab3:
    st.markdown(f'<div class="bar"><div class="bar-i">🛠</div><div><div class="bar-t">Other Equipment</div><div class="bar-s">lightweight extraction</div></div><div class="bar-r"></div><div class="bar-b">{len(oe)} rows</div></div>', unsafe_allow_html=True)
    if not oe:
        st.info("No other equipment data found.")
    else:
        oe_df = pd.DataFrame(oe)
        oe_df.columns = ['Section', 'Description', 'Last Date / O/H', 'Run Hrs']
        oe_df['Run Hrs'] = oe_df['Run Hrs'].apply(lambda x: str(x) if x else "—")
        
        st.dataframe(
            oe_df, use_container_width=True, hide_index=True,
            column_config={
                'Run Hrs': st.column_config.TextColumn("Run Hrs", width="small")
            }
        )

with tab4:
    issue_df = build_issue_df(issues)
    if issue_df.empty:
        st.success("No issues found.")
    else:
        st.dataframe(issue_df, use_container_width=True, hide_index=True)

with tab5:
    st.markdown("### Raw text preview")
    st.dataframe(pd.DataFrame({'line': debug['raw_lines_preview']}), use_container_width=True, hide_index=True, height=420)

    st.markdown("### Raw table previews")
    for tbl in debug['table_debug']:
        with st.expander(f"Table {tbl['table_index']} · {tbl['rows']} rows × {tbl['cols']} cols"):
            st.dataframe(pd.DataFrame(tbl['sample']), use_container_width=True, hide_index=True)

st.markdown("## Saved Reports")
hist = load_recent_reports(20)
if hist.empty:
    st.info("No saved reports yet.")
else:
    st.dataframe(hist, use_container_width=True, hide_index=True)
